from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Detection and Analysis of Pill', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Project Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.bold = True

# Add date
date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Blank line

# Add table of contents structure
doc.add_page_break()

# Chapter 1
doc.add_heading('Chapter 1: Introduction', 1)
doc.add_heading('1.1 Objective', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.2 Plan of Action', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.3 Literature Survey', 2)
doc.add_heading('1.3.1 Comparison Table', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.4 Business Context/Impact', 2)
doc.add_heading('1.4.1 Existing System Study', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.5 Proposed System', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.6 System Requirements', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.7 Functional and Non-functional Requirements', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('1.8 Feasibility Study', 2)
doc.add_paragraph('', style='List Bullet')

# Chapter 2
doc.add_page_break()
doc.add_heading('Chapter 2: Dataset Description', 1)
doc.add_heading('2.1 Data Sources', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('2.2 Data Description', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('2.3 Data Cleaning/Preprocessing', 2)
doc.add_heading('2.3.1 Handling missing values', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('2.3.2 Outlier detection and treatment', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('2.3.3 Data transformations', 3)
doc.add_paragraph('', style='List Bullet')

# Chapter 3
doc.add_page_break()
doc.add_heading('Chapter 3: Exploratory Data Analysis (EDA)', 1)
doc.add_heading('3.1 Key Insights', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('3.2 Visualizations', 2)
doc.add_heading('3.2.1 Correlations', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('3.2.2 Distribution plots', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('3.2.3 Anomaly detection', 3)
doc.add_paragraph('', style='List Bullet')

# Chapter 4
doc.add_page_break()
doc.add_heading('Chapter 4: Feature Engineering', 1)
doc.add_heading('4.1 Feature Selection/Extraction', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('4.2 New Features Created', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('4.3 Feature Scaling/Normalization', 2)
doc.add_paragraph('', style='List Bullet')

# Chapter 5
doc.add_page_break()
doc.add_heading('Chapter 5: Model Development', 1)
doc.add_heading('5.1 Algorithms Considered', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('5.2 Model Selection', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('5.3 Hyperparameter Tuning', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('5.4 Training Process', 2)
doc.add_heading('5.4.1 Train/test split', 3)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('5.4.2 Cross-validation details', 3)
doc.add_paragraph('', style='List Bullet')

# Chapter 6
doc.add_page_break()
doc.add_heading('Chapter 6: Model Evaluation', 1)
doc.add_heading('6.1 Evaluation Metrics', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('6.2 Performance Results', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('6.3 Model Comparison', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('6.4 Error Analysis', 2)
doc.add_paragraph('', style='List Bullet')

# Chapter 7
doc.add_page_break()
doc.add_heading('Chapter 7: Deployment Strategy', 1)
doc.add_heading('7.1 Model Serialization', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('7.2 Deployment Tools/Frameworks', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('7.3 Integration Process', 2)
doc.add_paragraph('', style='List Bullet')

# Chapter 8
doc.add_page_break()
doc.add_heading('Chapter 8: Monitoring and Maintenance', 1)
doc.add_heading('8.1 Performance Tracking', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('8.2 Retraining Strategy', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('8.3 Monitoring Tools', 2)
doc.add_paragraph('', style='List Bullet')

# Chapter 9
doc.add_page_break()
doc.add_heading('Chapter 9: Conclusion', 1)
doc.add_paragraph('', style='List Bullet')

# Chapter 10
doc.add_page_break()
doc.add_heading('Chapter 10: Future Scope of the Project', 1)
doc.add_paragraph('', style='List Bullet')

# Chapter 11
doc.add_page_break()
doc.add_heading('Chapter 11: Appendix', 1)
doc.add_heading('11.1 Code Repository', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('11.2 List of Abbreviations and Symbols Used', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('11.3 Publications Certificates/Awards', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('11.4 Published Paper', 2)
doc.add_paragraph('', style='List Bullet')
doc.add_heading('11.5 References/Bibliography', 2)
doc.add_paragraph('', style='List Bullet')

# Save the document
output_path = 'Detection_and_Analysis_of_Pill_Documentation.docx'
doc.save(output_path)
print(f"✓ Word document created successfully: {output_path}")
